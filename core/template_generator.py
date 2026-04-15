"""Generate certificate Word templates for each checklist item."""

import os
import re
import json
import zipfile
import io
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


DATA_DIR = os.path.join(os.path.dirname(__file__), "..", "data")


def generate_all_templates(
    checklist_items: list,
    psol_matches: dict,
    project_info: dict,
    output_dir: str,
    uploaded_reports: dict = None,
) -> list:
    """Generate certificate templates for all applicable checklist items.

    Args:
        checklist_items: Parsed checklist items (from checklist_parser or checklist_items.json)
        psol_matches: Dict mapping item id/name -> list of matched PSOLs
        project_info: {job_number, building, attention, company, address, date, ...}
        output_dir: Directory to save generated .docx files
        uploaded_reports: Dict mapping report type -> list of report detail dicts
                         e.g. {"fire": [{reference, company, date}], "access": [...]}

    Returns:
        List of generated file paths
    """
    os.makedirs(output_dir, exist_ok=True)
    generated = []
    uploaded_reports = uploaded_reports or {}

    for item in checklist_items:
        if _is_admin_item(item):
            continue

        item_id = item.get("id", item.get("name", ""))
        matched_psols = psol_matches.get(item_id, [])

        filepath = _generate_single_template(item, matched_psols, project_info, output_dir, uploaded_reports)
        if filepath:
            generated.append(filepath)

    return generated


def generate_zip(
    checklist_items: list,
    psol_matches: dict,
    project_info: dict,
    uploaded_reports: dict = None,
) -> bytes:
    """Generate all templates and return as a ZIP file bytes."""
    buf = io.BytesIO()
    uploaded_reports = uploaded_reports or {}
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for item in checklist_items:
            if _is_admin_item(item):
                continue

            item_id = item.get("id", item.get("name", ""))
            matched_psols = psol_matches.get(item_id, [])

            doc = _build_template(item, matched_psols, project_info, uploaded_reports)
            if doc:
                filename = _make_filename(item, project_info)
                doc_buf = io.BytesIO()
                doc.save(doc_buf)
                zf.writestr(filename, doc_buf.getvalue())

    return buf.getvalue()


def _is_admin_item(item: dict) -> bool:
    """Check if this is an administrative item that doesn't need a cert template."""
    skip_ids = {
        "cou_form", "government_levies", "trade_clearances",
        "survey_certificate", "fire_brigade_clearance",
        "developers_consent", "performance_solutions",
        "tccs_operational_acceptance", "access_consultant_signoff",
        "weatherproofing_facade", "fire_door_schedule",
    }
    item_id = item.get("id", "")
    if item_id in skip_ids:
        return True

    skip_names = {
        "COU Form", "Government Levies", "Trade Clearances",
        "Survey Certificate", "Fire Brigade Clearance",
        "Developers Consent", "Performance Solutions",
        "TCCS Operational Acceptance", "Access Consultant Sign Off",
        "Weatherproofing - Facade Engineer Sign Off", "Fire Door Schedule",
    }
    name = item.get("name", "")
    return name in skip_names


def _clean_address(project_info: dict) -> str:
    """Build a clean address field - just territory and building address.
    Strips out 'RE: Commercial Certificate Checklist...' text.
    """
    raw_address = project_info.get("address", "")
    building = project_info.get("building", "")

    # Strip out "RE: Commercial Certificate Checklist..." line
    lines = raw_address.split("\n")
    clean_lines = []
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Skip lines containing "RE:" or "Commercial Certificate Checklist"
        if re.match(r'^RE:\s', line, re.IGNORECASE):
            continue
        if "certificate checklist" in line.lower():
            continue
        clean_lines.append(line)

    if clean_lines:
        return "\n".join(clean_lines)

    # Fallback: just use building address
    return building


def _make_filename(item: dict, project_info: dict) -> str:
    """Generate filename for a certificate template."""
    name = item.get("name", item.get("id", "Certificate"))
    # Clean for filename
    name = name.replace("/", "-").replace("\\", "-").replace(":", "")
    name = name.replace("  ", " ").strip()
    job = project_info.get("job_number", "")
    # Avoid "Certificate Certificate"
    if "Certificate" in name or "Clearance" in name:
        return f"{job} - {name}.docx"
    return f"{job} - {name} Certificate.docx"


def _generate_single_template(
    item: dict,
    matched_psols: list,
    project_info: dict,
    output_dir: str,
    uploaded_reports: dict = None,
) -> str:
    """Generate a single certificate template .docx file."""
    doc = _build_template(item, matched_psols, project_info, uploaded_reports or {})
    if not doc:
        return None

    filename = _make_filename(item, project_info)
    filepath = os.path.join(output_dir, filename)
    doc.save(filepath)
    return filepath


def _build_template(item: dict, matched_psols: list, project_info: dict, uploaded_reports: dict = None) -> Document:
    """Build a certificate template Document matching the existing residential format."""
    uploaded_reports = uploaded_reports or {}
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- TITLE ---
    title_name = _get_title(item)
    title = doc.add_heading(title_name, level=2)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in title.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # spacer

    # --- PROJECT INFO TABLE ---
    project_name = project_info.get("building", project_info.get("description", ""))
    address = _clean_address(project_info)
    building_desc = project_info.get("building", "")

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    _set_cell(table.cell(0, 0), "Project Name", bold=True)
    _set_cell(table.cell(0, 1), project_name)
    _set_cell(table.cell(1, 0), "Address", bold=True)
    _set_cell(table.cell(1, 1), address)
    _set_cell(table.cell(2, 0), "Building to be certified", bold=True)
    _set_cell(table.cell(2, 1), building_desc)

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()  # spacer

    # --- CERTIFICATION STATEMENT ---
    p = doc.add_paragraph()
    run = p.add_run("I hereby certify that:")
    run.bold = True
    run.font.size = Pt(11)

    p2 = doc.add_paragraph()
    p2.add_run("The installation is in accordance with the nominated Standards of Performance.")

    doc.add_paragraph()  # spacer

    # --- STANDARDS TABLE ---
    standards_table = doc.add_table(rows=2, cols=2)
    standards_table.style = 'Table Grid'
    standards_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    _set_cell(standards_table.cell(0, 0), "Measure and/or system", bold=True)
    _set_cell(standards_table.cell(0, 1), "Standards of Performance", bold=True)

    # Content row
    item_name = item.get("name", "")
    _set_cell(standards_table.cell(1, 0), item_name)

    # Build standards text - includes report references for matching report types
    item_report_details = _get_item_report_details(item, uploaded_reports)
    standards_text = _build_standards_text(item, matched_psols, project_info, item_report_details)
    _set_cell_multiline(standards_table.cell(1, 1), standards_text)

    # Set column widths
    for row in standards_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()  # spacer

    # --- PERFORMANCE SOLUTION REQUIREMENTS (if specific PSOLs matched) ---
    if matched_psols:
        p = doc.add_paragraph()
        run = p.add_run("Performance Solution Requirements:")
        run.bold = True
        run.font.size = Pt(11)

        # Get full report details for header
        report_details = _get_report_details(matched_psols)
        if report_details:
            p2 = doc.add_paragraph()
            p2.add_run(f"The following Performance Solutions from the {', '.join(report_details)} apply to this certificate:")
            p2.paragraph_format.space_after = Pt(6)

        for psol in matched_psols:
            ps_num = psol["ps_number"]
            summary = psol.get("summary", "")
            clauses = ", ".join(psol.get("clauses", []))

            bullet = doc.add_paragraph(style='List Bullet')
            run = bullet.add_run(f"{ps_num}")
            run.bold = True
            bullet.add_run(f" ({clauses}): " if clauses else ": ")
            bullet.add_run(summary)
            bullet.paragraph_format.space_after = Pt(3)

        doc.add_paragraph()  # spacer

    doc.add_paragraph()  # spacer

    # --- QUALIFICATION STATEMENT ---
    p = doc.add_paragraph()
    p.add_run(
        "I am a properly qualified person and have a good working knowledge "
        "of the relevant codes and standards referenced above. "
        "(My qualifications and accreditations are listed below)"
    )

    doc.add_paragraph()  # spacer

    # --- SIGNOFF TABLE ---
    signoff_table = doc.add_table(rows=8, cols=2)
    signoff_table.style = 'Table Grid'
    signoff_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    signoff_fields = [
        "Relevant qualifications and accreditations",
        "Name",
        "Company",
        "ABN",
        "Address",
        "Phone No.",
        "Email",
    ]
    for i, field in enumerate(signoff_fields):
        _set_cell(signoff_table.cell(i, 0), field, bold=True)
        _set_cell(signoff_table.cell(i, 1), "")

    # Last row: Signature and Date
    _set_cell(signoff_table.cell(7, 0), "Signature", bold=True)
    _set_cell(signoff_table.cell(7, 1), "")

    for row in signoff_table.rows:
        row.cells[0].width = Cm(5)
        row.cells[1].width = Cm(11)

    doc.add_paragraph()  # spacer

    # Date field
    p = doc.add_paragraph()
    run = p.add_run("Date: ")
    run.bold = True
    p.add_run("____________________")

    # --- ADDITIONAL NOTES ---
    notes = _get_additional_notes(item)
    if notes:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run("Additional Notes:")
        run.bold = True
        run.font.size = Pt(11)
        for note in notes:
            np = doc.add_paragraph(style='List Bullet')
            np.add_run(note)

    return doc


def _get_title(item: dict) -> str:
    """Get the certificate title from the checklist item."""
    name = item.get("name", "").upper()

    # Clean up - don't double up on "CERTIFICATE"
    if "CERTIFICATE" in name or "CLEARANCE" in name:
        return name
    return f"{name} - INSTALLATION CERTIFICATE"


def _build_standards_text(item: dict, matched_psols: list, project_info: dict, item_report_details: list = None) -> list:
    """Build the standards of performance text as a list of lines."""
    lines = []
    item_report_details = item_report_details or []

    # NCC clauses - use year/amendment from project info if available
    ncc_clauses = item.get("ncc_clauses", [])
    if ncc_clauses:
        ncc_year = project_info.get("ncc_year", "2022")
        ncc_amendment = project_info.get("ncc_amendment", "")
        ncc_label = f"NCC {ncc_year}"
        if ncc_amendment:
            ncc_label += f" Amendment {ncc_amendment}"
        clause_str = ", ".join(ncc_clauses)
        lines.append(f"{ncc_label} Volume 1, {clause_str}")

    # Australian Standards
    standards = item.get("standards", [])
    for std in standards:
        lines.append(std)

    # Report references from item's report_types (e.g., FER for fire items)
    seen_report_refs = set()
    for detail in item_report_details:
        lines.append(f"and the approved {detail}")
        # Track the raw reference to avoid duplicating from PSOLs
        for word in ["Fire Engineering Report", "Access Solutions Report",
                      "Facade Engineer Report", "Waterproofing Report"]:
            detail_ref = detail.replace(word, "").strip()
            if detail_ref:
                seen_report_refs.add(detail_ref)

    # Additional report references from matched PSOLs (only if from a different report)
    if matched_psols:
        psol_details = _get_report_details(matched_psols)
        for detail in psol_details:
            # Check if this report reference is already covered by report_types
            detail_ref = detail.replace("Performance Solution Report", "").strip()
            if detail_ref not in seen_report_refs:
                lines.append(f"and the approved {detail}")

    # If no clauses or standards, use the requirement text from checklist
    if not lines:
        req = item.get("requirement", item.get("description", ""))
        if req:
            lines.append(req)

    return lines


def _get_item_report_details(item: dict, uploaded_reports: dict) -> list:
    """Get report detail strings for an item based on its report_types.

    Args:
        item: Checklist item with report_types field
        uploaded_reports: Dict mapping report type -> list of report dicts
                         e.g. {"fire": [{"reference": "...", "company": "...", "date": "..."}]}

    Returns:
        List of detail strings like "Fire Engineering Report 20240243 FER V2 prepared by BSE dated 17 Dec 2024"
    """
    details = []
    report_types = item.get("report_types", [])

    report_type_labels = {
        "fire": "Fire Engineering Report",
        "access": "Access Solutions Report",
        "facade": "Facade Engineer Report",
        "waterproofing": "Waterproofing Report",
    }

    for rtype in report_types:
        reports = uploaded_reports.get(rtype, [])
        for report in reports:
            ref = report.get("reference", "")
            company = report.get("company", "")
            date = report.get("date", "")

            label = report_type_labels.get(rtype, "Report")
            parts = [label]
            if ref:
                parts[0] = f"{label} {ref}"
            if company:
                parts.append(f"prepared by {company}")
            if date:
                parts.append(f"dated {date}")

            detail = " ".join(parts)
            if detail not in details:
                details.append(detail)

    return details


def _set_cell(cell, text: str, bold: bool = False):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True


def _set_cell_multiline(cell, lines: list):
    """Set cell with multiple lines of text."""
    cell.text = ""
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(line)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'


def _get_report_details(matched_psols: list) -> list:
    """Build full report detail strings from matched PSOLs.

    Returns list like:
        ["Performance Solution Report ACT22080 FER Rev E prepared by Lit Consulting Pty Ltd dated 19 February 2026"]
    """
    # Group by report reference to avoid duplicates
    seen = set()
    details = []

    for psol in matched_psols:
        ref = psol.get("report_reference", "")
        if not ref or ref in seen:
            continue
        seen.add(ref)

        company = psol.get("report_company", "")
        date = psol.get("report_date", "")

        parts = [f"Performance Solution Report {ref}"]
        if company:
            parts.append(f"prepared by {company}")
        if date:
            parts.append(f"dated {date}")

        details.append(" ".join(parts))

    return details


def _get_additional_notes(item: dict) -> list:
    """Get any additional notes specific to certain certificate types."""
    notes_map = {
        "fire_rated_penetrations": [
            "A schedule of ALL fire rated penetrations must be provided",
            "Schedule must cover all service trades: Electrical, Fire, Hydraulic, Mechanical, Communications",
        ],
        "lightweight_fire_rating": [
            "Must identify the specific system installed",
            "Must state the minimum FRL which the system achieves",
        ],
        "glazing": [
            "Must itemize the type of glazing for each location",
            "Must state the applicable section of AS 1288 to which each element complies",
        ],
        "waterproofing_internal": [
            "Must identify the waterproofing product used",
        ],
        "waterproofing_external": [
            "Must identify the waterproofing product used",
        ],
        "flooring": [
            "Must identify the specific floor covering products used",
        ],
        "sarking": [
            "Must identify the sarking product used",
            "Must verify the R-value achieved",
        ],
    }
    return notes_map.get(item.get("id", ""), [])
