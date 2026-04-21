"""Extract Performance Solutions from PDF reports (FER, Access, Facade, etc.)."""

import re
import fitz  # PyMuPDF
from .checklist_parser import extract_ncc_clauses


def extract_psols_from_pdf(pdf_path: str) -> dict:
    """Extract Performance Solutions from a report (PDF or DOCX).

    Returns dict with:
        report: {filename, title, author, reference}
        psols: [{ps_number, clauses, summary}]
    """
    filename = pdf_path.replace("\\", "/").split("/")[-1]
    metadata = {
        "filename": filename,
        "title": "",
        "author": "",
        "reference": "",
        "company": "",
        "date": "",
    }

    # Handle DOCX files too
    if pdf_path.lower().endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(pdf_path)
            all_text = "\n".join(p.text for p in doc.paragraphs)
            # Also include text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        all_text += "\n" + cell.text
        except Exception as e:
            return {"report": metadata, "psols": []}
    else:
        doc = fitz.open(pdf_path)
        metadata["title"] = doc.metadata.get("title", "")
        metadata["author"] = doc.metadata.get("author", "")

        # Extract text from all pages
        all_text = ""
        for page in doc:
            all_text += page.get_text() + "\n"
        doc.close()

    # Extract report metadata from first few pages
    first_pages = all_text[:5000]

    # Report reference - try multiple patterns
    # Pattern 1: "Document reference\nACT22080 FER Rev E" (may span lines)
    ref_match = re.search(r'Document\s*\n?\s*reference\s*\n\s*(.+?)(?:\n|$)', first_pages)
    if ref_match:
        metadata["reference"] = ref_match.group(1).strip()
    else:
        # Pattern 2: "ACT25036 FER Rev B | 19 March 2026" on title page
        ref_match = re.search(r'([A-Z]{2,}\d+\s+FER\s+Rev\s+[A-Z0-9]+)', first_pages)
        if ref_match:
            metadata["reference"] = ref_match.group(1).strip()
        else:
            # Pattern 3: "Document Issue\n2" + "Project Number\n20240243" -> build reference
            issue_match = re.search(r'Document Issue\s*\n\s*(\S+)', first_pages)
            proj_match = re.search(r'Project Number\s*\n\s*(\S+)', first_pages)
            if issue_match or proj_match:
                parts = []
                if proj_match:
                    parts.append(proj_match.group(1).strip())
                type_match = re.search(r'(Fire Engineering)\s*\n?\s*Report', first_pages)
                if type_match:
                    parts.append("FER")
                if issue_match:
                    parts.append(f"V{issue_match.group(1).strip()}")
                if parts:
                    metadata["reference"] = " ".join(parts)

    # Company name - try multiple patterns
    # Pattern 1: "Pty Ltd" in name
    company_match = re.search(r'^(.+?Pty\s+Ltd.*?)$', first_pages, re.MULTILINE)
    if company_match:
        metadata["company"] = company_match.group(1).strip()
    else:
        # Pattern 2: "Prepared By:\n\nCompany Name" or company header like "BSE | Building Services Engineers"
        prep_match = re.search(r'Prepared By:\s*\n\s*\n?\s*(.+?)(?:\n|$)', first_pages)
        if prep_match:
            metadata["company"] = prep_match.group(1).strip()

    # Report date - try multiple patterns
    # Pattern 1: "Rev E | 19 February 2026"
    date_match = re.search(r'Rev\s+[A-Z]\s*\|\s*(\d{1,2}\s+\w+\s+\d{4})', first_pages)
    if date_match:
        metadata["date"] = date_match.group(1).strip()
    else:
        # Pattern 2: "Document Date\n17 December 2024"
        date_match = re.search(r'Document Date\s*\n\s*(\d{1,2}\s+\w+\s+\d{4})', first_pages)
        if date_match:
            metadata["date"] = date_match.group(1).strip()
        else:
            # Pattern 3: revision history dates (DD/MM/YYYY)
            rev_dates = re.findall(r'(\d{2}/\d{2}/\d{4})', first_pages)
            if rev_dates:
                metadata["date"] = rev_dates[-1]

    # Extract PSOLs using multiple strategies
    psols = _extract_from_summary_table(all_text)

    if not psols:
        # Fallback: try PS-prefix pattern
        psols = _extract_ps_prefix_pattern(all_text)

    if not psols:
        # Fallback: try section headings pattern
        psols = _extract_from_section_headings(all_text)

    return {"report": metadata, "psols": psols}


def extract_psols_from_multiple(pdf_paths: list) -> list:
    """Extract PSOLs from multiple PDF reports and combine results."""
    all_psols = []
    for path in pdf_paths:
        result = extract_psols_from_pdf(path)
        report = result["report"]
        for psol in result["psols"]:
            psol["report_filename"] = report["filename"]
            psol["report_reference"] = report["reference"]
            psol["report_company"] = report["company"]
            psol["report_date"] = report["date"]
            all_psols.append(psol)

    # Sort combined results
    all_psols.sort(key=lambda x: _ps_sort_key(x["ps_number"]))
    return all_psols


def _extract_from_summary_table(text: str) -> list:
    """Extract PSOLs from numbered summary tables in FER format.

    The FER uses a table like:
    No  Description of Issue  BCA DtS Provision  Performance Requirements
    1.  The fire resistance...  Clause C2D2...     C1P1 and C1P2
    """
    psols = []
    seen = set()

    # Find the summary table region - look for "Summary of performance solutions"
    # or "Description of Issue" headers
    table_start = None
    for marker in [
        "Summary of performance solutions",
        "Description of Issue",
        "Summary of Performance Solutions",
    ]:
        idx = text.find(marker)
        if idx != -1:
            table_start = idx
            break

    if table_start is None:
        return []

    # Work through the text after the table header
    table_text = text[table_start:]

    # Find where the table ends (look for section breaks)
    table_end_markers = [
        "The fire engineering performance solutions documented",
        "Contents\n",
        "\nContents ",
        "subject to the implementation",
    ]
    table_end = len(table_text)
    for marker in table_end_markers:
        idx = table_text.find(marker)
        if idx != -1 and idx < table_end:
            table_end = idx

    table_text = table_text[:table_end]

    # Parse numbered entries: "N." or "N. " at start of line followed by description
    # Then look for "Clause" references nearby
    # FER format splits across lines, so we need to be flexible

    # Strategy: find all numbered items and their associated clause text
    # Pattern: number followed by period, then text until next number or end
    entries = _split_numbered_entries(table_text)

    for num, entry_text in entries:
        ps_number = f"PS{num}"

        if ps_number in seen:
            continue

        # Check if removed
        if re.search(r'has been removed|REMOVED', entry_text, re.IGNORECASE):
            continue

        # Extract description (first part before clause references)
        description = _clean_description(entry_text)

        # Extract NCC clauses
        clauses = extract_ncc_clauses(entry_text)

        # Also look for "Clause X" patterns specific to FER format
        for m in re.finditer(r'Clauses?\s+([A-Z]\d[A-Z]\d+(?:,\s*[A-Z]\d[A-Z]\d+)*)', entry_text):
            for clause in re.findall(r'[A-Z]\d[A-Z]\d+', m.group(0)):
                if clause not in clauses:
                    clauses.append(clause)
            clauses = sorted(set(clauses))

        if description:
            seen.add(ps_number)
            psols.append({
                "ps_number": ps_number,
                "clauses": clauses,
                "summary": description,
            })

    return psols


def _split_numbered_entries(text: str) -> list:
    """Split text into numbered entries like '1. description...' '2. description...'"""
    entries = []

    # Find all positions where a numbered entry starts
    # Pattern: start of line or after whitespace, digit(s), period, space
    positions = []
    for m in re.finditer(r'(?:^|\n)\s*(\d{1,2})\.\s+', text):
        num = int(m.group(1))
        if 1 <= num <= 99:
            positions.append((num, m.start(), m.end()))

    # Extract text between positions
    for i, (num, start, content_start) in enumerate(positions):
        if i + 1 < len(positions):
            end = positions[i + 1][1]
        else:
            end = len(text)
        entry_text = text[content_start:end].strip()
        entries.append((num, entry_text))

    return entries


def _clean_description(entry_text: str) -> str:
    """Extract clean description from an entry, removing clause/requirement refs."""
    # The entry typically has: description text, then BCA DtS clause, then performance reqs
    # Try to get just the description part

    # Remove performance requirement codes (C1P1, D1P4, E1P3, etc.)
    text = re.sub(r'\b[A-Z]\dP\d+\b', '', entry_text)

    # Remove "and" connecting perf requirements
    text = re.sub(r'\s+and\s+$', '', text.strip())

    # Remove clause prefix words at the end
    text = re.sub(r'\s*Clauses?\s+[A-Z]\d.*$', '', text, flags=re.DOTALL)

    # Remove specification references at the end
    text = re.sub(r'\s*(?:and\s+)?specification\s+\d+\s*$', '', text, flags=re.IGNORECASE)

    # Clean whitespace
    text = re.sub(r'\s+', ' ', text).strip()

    # Limit length
    if len(text) > 400:
        text = text[:400].rsplit(' ', 1)[0] + "..."

    return text


def _extract_ps_prefix_pattern(text: str) -> list:
    """Extract PSOLs that use PS prefix format (PS1, PS-A, etc.)."""
    psols = []
    seen = set()

    for match in re.finditer(r'\b(PS[-\s]?(?:\d+|[A-Z]))\b', text):
        ps_number = match.group(1).replace(" ", "")
        if ps_number in seen:
            continue
        seen.add(ps_number)

        start = max(0, match.start() - 200)
        end = min(len(text), match.end() + 500)
        context = text[start:end]
        clauses = extract_ncc_clauses(context)

        # Get surrounding text for summary
        line_start = text.rfind("\n", 0, match.start())
        line_start = 0 if line_start == -1 else line_start + 1
        line_end = text.find("\n\n", match.end())
        if line_end == -1 or line_end > match.end() + 300:
            line_end = match.end() + 300
        summary = re.sub(r'\s+', ' ', text[line_start:line_end]).strip()[:400]

        if "REMOVED" in summary.upper():
            continue

        psols.append({"ps_number": ps_number, "clauses": clauses, "summary": summary})

    return psols


def _extract_from_section_headings(text: str) -> list:
    """Extract PSOLs from section heading format: 'Performance solution N – Title'."""
    psols = []
    seen = set()

    pattern = r'[Pp]erformance [Ss]olution\s+(\d+)\s*[-–]\s*(.+?)(?:\n|$)'
    for match in re.finditer(pattern, text):
        num = int(match.group(1))
        title = match.group(2).strip()
        ps_number = f"PS{num}"

        if ps_number in seen:
            continue
        if "removed" in title.lower():
            continue
        seen.add(ps_number)

        # Get context for clauses
        start = match.start()
        end = min(len(text), match.end() + 1000)
        context = text[start:end]
        clauses = extract_ncc_clauses(context)

        psols.append({"ps_number": ps_number, "clauses": clauses, "summary": title})

    return psols


def _ps_sort_key(ps_number: str) -> tuple:
    """Sort key: PS1-PS99, then PS-A, PS-B, etc."""
    match = re.match(r'PS-?(\d+)', ps_number)
    if match:
        return (0, int(match.group(1)))
    match = re.match(r'PS-?([A-Z])', ps_number)
    if match:
        return (1, ord(match.group(1)))
    return (2, 0)


def detect_report_type(pdf_path: str, filename: str = "") -> str:
    """Detect the type of performance solution report from its content and filename.

    Returns one of: "fire", "access", "facade", "waterproofing", "section_j",
    "acoustic", "mechanical", "bushfire", or "" if unknown.
    """
    filename_lower = filename.lower()

    # Check filename first for quick classification
    if any(k in filename_lower for k in ["fer", "fire engineer", "fire engineering"]):
        return "fire"
    if any(k in filename_lower for k in ["access", "das", "disability"]):
        return "access"
    if any(k in filename_lower for k in ["facade", "façade", "cladding"]):
        return "facade"
    if any(k in filename_lower for k in ["waterproof", "water proof", "f3p1"]):
        return "waterproofing"
    if any(k in filename_lower for k in ["section j", "sectionj", "section-j", "energy efficiency", "energy report", "jv3", "basix"]):
        return "section_j"
    if any(k in filename_lower for k in ["acoustic"]):
        return "acoustic"
    if any(k in filename_lower for k in ["mechanical", "hvac"]):
        return "mechanical"
    if any(k in filename_lower for k in ["bushfire", "bal "]):
        return "bushfire"

    # Read first few pages to detect from content
    try:
        if pdf_path.lower().endswith(".docx"):
            # Handle docx for report type detection
            from docx import Document as DocxDocument
            doc = DocxDocument(pdf_path)
            first_pages = "\n".join(p.text for p in doc.paragraphs[:200])
        else:
            doc = fitz.open(pdf_path)
            first_pages = ""
            for i in range(min(5, doc.page_count)):
                first_pages += doc[i].get_text() + "\n"
            doc.close()
    except Exception:
        return ""

    first_lower = first_pages.lower()

    if any(k in first_lower for k in ["fire engineering report", "fire engineering\nreport", "fire engineer"]):
        return "fire"
    if any(k in first_lower for k in ["access report", "access solution", "access consultant", "disability"]):
        return "access"
    if any(k in first_lower for k in ["facade engineer", "façade engineer", "facade report", "cladding"]):
        return "facade"
    if any(k in first_lower for k in ["waterproofing report", "waterproofing solution", "f3p1"]):
        return "waterproofing"
    if any(k in first_lower for k in ["section j", "energy efficiency", "jv3 ", "basix certificate"]):
        return "section_j"
    if any(k in first_lower for k in ["acoustic report", "acoustic consultant"]):
        return "acoustic"
    if any(k in first_lower for k in ["mechanical services report", "hvac design"]):
        return "mechanical"
    if any(k in first_lower for k in ["bushfire attack level", "bushfire assessment"]):
        return "bushfire"

    # If it has performance solutions with fire-related clauses, assume fire
    if any(k in first_lower for k in ["performance solution", "bca dts clause"]):
        if re.search(r'\b[CE]\d[DP]\d', first_pages):
            return "fire"

    return ""


def match_psols_to_checklist(psols: list, checklist_items: list) -> dict:
    """Match extracted PSOLs to checklist items by NCC clause overlap.

    Returns dict mapping checklist item id/name -> list of matching PSOLs
    """
    matches = {}

    for item in checklist_items:
        item_id = item.get("id", item.get("name", ""))
        item_clauses = set(item.get("ncc_clauses", []))
        if not item_clauses:
            matches[item_id] = []
            continue

        matched = []
        for psol in psols:
            psol_clauses = set(psol.get("clauses", []))
            # Direct overlap
            if item_clauses & psol_clauses:
                if psol not in matched:
                    matched.append(psol)
                continue
            # Partial matches (Specification 20 <-> S20C4, etc.)
            for ic in item_clauses:
                for pc in psol_clauses:
                    if ic in pc or pc in ic:
                        if psol not in matched:
                            matched.append(psol)

        matches[item_id] = matched

    return matches
