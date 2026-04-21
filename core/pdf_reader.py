"""Extract text from PDF and DOCX files."""

import fitz  # PyMuPDF


def extract_text(file_path: str) -> str:
    """Extract all text from a PDF or DOCX file."""
    if file_path.lower().endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            parts = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        except Exception:
            return ""

    doc = fitz.open(file_path)
    text_parts = []
    for page in doc:
        text_parts.append(page.get_text())
    doc.close()
    return "\n".join(text_parts)


def extract_text_by_page(pdf_path: str) -> list:
    """Extract text from each page of a PDF, returned as a list."""
    doc = fitz.open(pdf_path)
    pages = []
    for page in doc:
        pages.append(page.get_text())
    doc.close()
    return pages


def extract_report_metadata(pdf_path: str) -> dict:
    """Try to extract report metadata (title, author, revision) from first pages."""
    doc = fitz.open(pdf_path)
    metadata = {
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "filename": pdf_path.split("/")[-1].split("\\")[-1],
    }
    doc.close()
    return metadata
