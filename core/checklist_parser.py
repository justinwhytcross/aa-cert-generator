"""Parse COU checklist .docx files to extract project details and checklist items."""

import re
from docx import Document


def parse_checklist(docx_path: str) -> dict:
    """Parse a COU checklist .docx and return structured data.

    Returns dict with:
        project: {job_number, date, attention, company, address, description, building}
        items: [{number, name, requirement, status, ncc_clauses, standards}]
    """
    doc = Document(docx_path)

    project = _extract_project_info(doc)
    items = _extract_checklist_items(doc)

    return {"project": project, "items": items}


def _extract_project_info(doc: Document) -> dict:
    """Extract project header info from the paragraphs before the table."""
    info = {
        "job_number": "",
        "date": "",
        "attention": "",
        "company": "",
        "address": "",
        "description": "",
        "building": "",
    }

    paragraphs_text = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs_text.append(text)

    full_text = "\n".join(paragraphs_text)

    # Extract job number
    job_match = re.search(r'Job\s*No\s*[:\.]?\s*(\d+)', full_text)
    if job_match:
        info["job_number"] = job_match.group(1)

    # Extract date (first line is typically the date)
    if paragraphs_text:
        date_match = re.match(r'^[A-Za-z]+\s+\d{1,2}\s+[A-Za-z]+\s+\d{4}', paragraphs_text[0])
        if date_match:
            info["date"] = date_match.group(0)

    # Extract attention
    att_match = re.search(r'Attention\s*[:\.]?\s*(.+?)(?:\n|$)', full_text)
    if att_match:
        info["attention"] = att_match.group(1).strip()

    # Extract RE line (project description + building)
    re_match = re.search(r'RE\s*[:\.]?\s*Commercial Certificate Checklist.*?[-\u2013]\s*(.+?)(?:\n|$)', full_text)
    if re_match:
        info["building"] = re_match.group(1).strip()

    # Extract NCC year and amendment from checklist header
    # e.g., "NCC2022" or "NCC2022 (1)" where (1) = Amendment 1
    ncc_match = re.search(r'NCC\s*(\d{4})\s*(?:\((\d+)\))?', full_text)
    if ncc_match:
        ncc_year = ncc_match.group(1)
        amendment = ncc_match.group(2) if ncc_match.group(2) else ""
        info["ncc_year"] = ncc_year
        info["ncc_amendment"] = amendment

    # Extract description of work
    desc_match = re.search(r'Description of work\s*[:\.]?\s*(.+?)(?:\n|$)', full_text)
    if desc_match:
        info["description"] = desc_match.group(1).strip()

    # Extract company and address from paragraphs after Attention
    for i, text in enumerate(paragraphs_text):
        if "Attention" in text:
            # Next paragraphs are typically company, street address, city
            remaining = paragraphs_text[i+1:i+4]
            if len(remaining) >= 1:
                info["company"] = remaining[0]
            if len(remaining) >= 2:
                info["address"] = "\n".join(remaining[1:])
            break

    return info


def _extract_checklist_items(doc: Document) -> list:
    """Extract checklist items from the main table."""
    items = []

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 3:
            continue

        # Check if this is the checklist table (look for header row with #, Item, etc.)
        header_text = " ".join(cell.text.strip() for cell in rows[0].cells)
        if "#" not in header_text and "Item" not in header_text:
            # Try second row (first might be a section header)
            if len(rows) > 1:
                header_text = " ".join(cell.text.strip() for cell in rows[1].cells)
            if "#" not in header_text and "Item" not in header_text:
                continue

        # Find header row index
        header_idx = 0
        for i, row in enumerate(rows):
            cells_text = [cell.text.strip() for cell in row.cells]
            if "#" in cells_text or "Item" in cells_text:
                header_idx = i
                break

        # Parse data rows
        for row in rows[header_idx + 1:]:
            cells = row.cells
            if len(cells) < 3:
                continue

            # Get cell texts
            cell_texts = [cell.text.strip() for cell in cells]

            # Skip section header rows (like "Certificates" or "Certificate Checklist Requirements")
            if len(cells) >= 4:
                number = cell_texts[0]
                name = cell_texts[1]
                requirement = cell_texts[2]
                status = cell_texts[3] if len(cells) > 3 else ""
            else:
                continue

            # Skip empty rows or header/note rows
            if not name or name in ("Item", "Certificates", "Certificate Checklist Requirements"):
                continue
            if name.startswith("Please note"):
                continue

            # Extract NCC clauses from requirement text
            ncc_clauses = extract_ncc_clauses(requirement)

            # Extract AS standards from requirement text
            standards = extract_standards(requirement)

            items.append({
                "number": number,
                "name": name,
                "requirement": requirement,
                "status": status,
                "ncc_clauses": ncc_clauses,
                "standards": standards,
            })

    return items


def extract_ncc_clauses(text: str) -> list:
    """Extract NCC clause references from text.

    Patterns: E1D2, C4D15, D3D17, Specification 7, Specification 20,
    Section B, Part F7, S20C4, F3D3, F8D5, C2D10, C2D11
    """
    clauses = set()

    # Standard clause pattern: letter-digit-letter-digit(s) e.g. E1D2, C4D15, D3D17
    for match in re.finditer(r'\b([A-Z]\d[A-Z]\d+)\b', text):
        clauses.add(match.group(1))

    # Specification pattern
    for match in re.finditer(r'Specification\s+(\d+)', text):
        clauses.add(f"Specification {match.group(1)}")

    # Spec short pattern: S20C4
    for match in re.finditer(r'\b(S\d+C\d+)\b', text):
        clauses.add(match.group(1))

    # Section pattern
    for match in re.finditer(r'Section\s+([A-Z])\b', text):
        clauses.add(f"Section {match.group(1)}")

    # Part pattern: Part F7
    for match in re.finditer(r'Part\s+([A-Z]\d+)', text):
        clauses.add(match.group(1))

    # Clause pattern with prefix: Clause E1D2, Clauses E4D5
    for match in re.finditer(r'Clauses?\s+([A-Z]\d[A-Z]\d+)', text):
        clauses.add(match.group(1))

    return sorted(clauses)


def extract_standards(text: str) -> list:
    """Extract Australian Standard references from text.

    Patterns: AS 2419.1 - 2021, AS/NZS 2293.1 - 2018, AS 3600:2018, AS4440 - 2004
    """
    standards = set()

    # AS or AS/NZS followed by number, optional dot+number, then year
    pattern = r'(AS(?:/NZS)?\s*\d+(?:\.\d+)?)\s*[-:]\s*(\d{4})'
    for match in re.finditer(pattern, text):
        std = f"{match.group(1).strip()} - {match.group(2)}"
        # Normalize spacing
        std = re.sub(r'\s+', ' ', std)
        standards.add(std)

    return sorted(standards)
